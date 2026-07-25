from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "employee-workspace-user-guide-ar.docx"
SCREENSHOTS = ROOT / "docs" / "employee-workspace-screenshots"
ENTRY_CARD_SCREENSHOT = Path(
    r"C:\Users\khamis\AppData\Local\Temp\codex-clipboard-4955a961-625f-4b8e-9e2e-e736fe8f82c7.png"
)

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 111, 132)
BORDER = "D9E2EC"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
SUCCESS_FILL = "E8F7EF"
CAUTION_FILL = "FFF4DA"
RISK_FILL = "FDECEC"
WHITE = "FFFFFF"


def set_document_rtl(doc: Document) -> None:
    settings = doc.settings._element
    if settings.find(qn("w:themeFontLang")) is None:
        lang = OxmlElement("w:themeFontLang")
        lang.set(qn("w:val"), "ar-SA")
        lang.set(qn("w:bidi"), "ar-SA")
        settings.append(lang)


def add_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, *, after=6, before=0, line=1.25, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    add_bidi(paragraph)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_para(doc, text: str, *, size=11, bold=False, color=NAVY, after=6, before=0, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph()
    style_paragraph(p, after=after, before=before, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text: str, level: int = 1):
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, DARK_BLUE, 10, 5
    return add_para(doc, text, size=size, bold=True, color=color, before=before, after=after)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_setup(table, widths: list[int], fill_header=True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    if tbl_pr.find(qn("w:bidiVisual")) is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            if fill_header and row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for p in cell.paragraphs:
                style_paragraph(p, after=0)


def set_table_text(cell, text: str, *, bold=False, fill: str | None = None, color=NAVY, size=10.5) -> None:
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    style_paragraph(p, after=0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_label_detail_table(doc, rows: list[tuple[str, str]], *, title: str | None = None) -> None:
    if title:
        add_para(doc, title, size=10.5, bold=True, color=MUTED, before=4, after=4)
    table = doc.add_table(rows=len(rows), cols=2)
    widths = [2700, 6660]
    table_setup(table, widths, fill_header=False)
    for idx, (label, detail) in enumerate(rows):
        set_table_text(table.cell(idx, 0), label, bold=True, fill=LIGHT_FILL, size=10.2)
        set_table_text(table.cell(idx, 1), detail, size=10.2)
    add_para(doc, "", after=2)


def add_matrix_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table_setup(table, widths, fill_header=True)
    for idx, header in enumerate(headers):
        set_table_text(table.cell(0, idx), header, bold=True, size=10.2)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            set_table_text(table.cell(r_idx, c_idx), text, size=9.6)
    add_para(doc, "", after=2)


def add_bullet(doc, text: str, *, fill: str | None = None) -> None:
    if fill:
        table = doc.add_table(rows=1, cols=1)
        table_setup(table, [9360], fill_header=False)
        set_table_text(table.cell(0, 0), "• " + text, fill=fill, size=10.3)
        return
    p = doc.add_paragraph()
    style_paragraph(p, after=4)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.left_indent = Inches(0.375)
    run = p.add_run("• " + text)
    set_run_font(run, size=10.8, color=NAVY)


def add_figure(doc, path: Path, title: str, caption: str, *, width_in=6.2) -> None:
    if not path.exists():
        return
    add_para(doc, title, size=10.5, bold=True, color=DARK_BLUE, before=8, after=4)
    p = doc.add_paragraph()
    style_paragraph(p, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_para(doc, caption, size=9.3, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_number(doc, number: int, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, after=4)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.left_indent = Inches(0.375)
    run = p.add_run(f"{number}. {text}")
    set_run_font(run, size=10.8, color=NAVY)


def configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    header_p = section.header.paragraphs[0]
    style_paragraph(header_p, after=0)
    run = header_p.add_run("Fleetify | دليل مساحة عمل الموظف")
    set_run_font(run, size=9.5, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    style_paragraph(footer_p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = footer_p.add_run("Fleetify - للاستخدام الداخلي")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    add_para(doc, "دليل تشغيلي للموظفين", size=11, bold=True, color=BLUE, after=4)
    add_para(doc, "مساحة عمل الموظف", size=29, bold=True, color=NAVY, after=6)
    add_para(
        doc,
        "شرح عملي لاستخدام صفحة employee-workspace في متابعة العقود، التحصيل، المهام، والتصعيد للشؤون القانونية.",
        size=13,
        color=MUTED,
        after=18,
    )
    add_label_detail_table(
        doc,
        [
            ("النظام", "Fleetify - شركة العراف لتأجير السيارات"),
            ("الجمهور", "موظفو التحصيل والمتابعة وخدمة العملاء"),
            ("الغرض", "توحيد طريقة العمل اليومية داخل مساحة العمل الخاصة بكل موظف"),
            ("الإصدار", "نسخة تدريب داخلية بالصور من النظام الحي - 25 يوليو 2026"),
        ],
    )
    add_bullet(
        doc,
        "هذا الدليل يركز على العمل اليومي داخل صفحة مساحة العمل. أي إعادة تعيين للعقود أو إلغاء تعيين تتم من إدارة الفريق حسب الصلاحيات المعتمدة.",
        fill=CAUTION_FILL,
    )
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def build_doc() -> None:
    doc = Document()
    set_document_rtl(doc)
    configure_sections(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "Calibri")

    add_cover(doc)

    add_heading(doc, "1. الهدف من الصفحة")
    add_para(
        doc,
        "مساحة عمل الموظف هي شاشة العمل اليومية في النظام الحي على alaraf.online، وتجمع أهم ما يحتاجه الموظف لمتابعة العملاء والعقود: المستحقات، المهام، الاتصالات، الدفعات، الملاحظات، والتصعيد القانوني عند الحاجة.",
    )
    add_bullet(doc, "بدل التنقل بين عدة أقسام، يبدأ الموظف من هذه الصفحة ويعالج الأولويات حسب الحالة والمبالغ والمهام.")
    add_bullet(doc, "الصفحة تعرض العقود المعيّنة للموظف، وتساعده على توثيق كل إجراء يتم على العميل.")
    add_bullet(doc, "كل إجراء مهم يجب أن يُسجل داخل النظام حتى يظهر في المتابعة والتقارير.")

    add_heading(doc, "2. طريقة الدخول إلى مساحة عملي")
    add_para(
        doc,
        "بعد تسجيل الدخول للنظام، تظهر في لوحة التحكم بطاقة باسم الفريق والمتابعة. من هذه البطاقة يضغط الموظف على زر مساحة عملي للدخول إلى صفحة العمل اليومية الخاصة به.",
    )
    add_bullet(doc, "زر مساحة عملي يفتح صفحة متابعة التحصيل والعقود والمهام الخاصة بالموظف.")
    add_bullet(doc, "زر إدارة الفريق مخصص للمديرين أو المسؤولين عن توزيع العقود، وليس هو مسار العمل اليومي للموظف.")
    add_figure(
        doc,
        ENTRY_CARD_SCREENSHOT,
        "صورة 1: بطاقة الفريق والمتابعة في لوحة التحكم",
        "يتم الدخول إلى مساحة العمل من زر مساحة عملي داخل هذه البطاقة.",
        width_in=3.2,
    )

    add_heading(doc, "3. نظرة سريعة على أجزاء الشاشة")
    add_matrix_table(
        doc,
        ["الجزء", "ماذا يعرض؟", "متى تستخدمه؟"],
        [
            ["الشريط العلوي", "اسم الصفحة، أزرار التحديث، التصدير، والتنبيهات", "في بداية اليوم أو عند تحديث البيانات أو استخراج تقرير"],
            ["بطاقات المؤشرات", "إجمالي العقود، المبالغ المستحقة، مهام اليوم، نقاط الأداء", "لمعرفة وضعك اليومي بسرعة"],
            ["الإجراءات السريعة", "تسجيل مكالمة، تسجيل دفعة، جدولة موعد، ملاحظة جديدة", "لتوثيق إجراء مباشر بدون مغادرة الصفحة"],
            ["التبويبات", "نظرة عامة، التحصيل الشهري، العقود، المهام", "للتنقل بين نوع العمل المطلوب"],
        ],
        [1900, 3860, 3600],
    )
    add_figure(
        doc,
        SCREENSHOTS / "01-overview-desktop.png",
        "صورة 2: الواجهة العامة لمساحة العمل من النظام الحي",
        "توضح الصورة الشريط العلوي، أزرار الإجراءات السريعة، بطاقات المؤشرات، والتبويبات الأساسية بعد تسجيل الدخول للنظام الحي.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "02-overview-mobile.png",
        "صورة 3: عرض الصفحة على الجوال",
        "في الجوال تظهر العناصر بشكل رأسي؛ استخدم التمرير للأسفل للوصول إلى المؤشرات والتبويبات وبطاقات المهام.",
        width_in=2.25,
    )
    add_figure(
        doc,
        SCREENSHOTS / "07-quick-actions-row.png",
        "صورة 4: أزرار الإجراءات السريعة",
        "هذه الأزرار هي أكثر أدوات الموظف استخدامًا خلال اليوم: مكالمة، دفعة، موعد، وملاحظة.",
    )

    add_heading(doc, "4. طريقة العمل اليومية المقترحة")
    daily_steps = [
        "افتح صفحة مساحة عملي وتأكد أن البيانات ظهرت بعد تسجيل الدخول.",
        "راجع بطاقات المؤشرات لمعرفة عدد العقود، المبالغ المستحقة، والمهام المطلوبة اليوم.",
        "ابدأ من تبويب نظرة عامة وراجع العملاء الذين يحتاجون اهتمامًا فوريًا.",
        "نفّذ الاتصال أو المتابعة، ثم سجّل النتيجة مباشرة باستخدام تسجيل مكالمة أو ملاحظة جديدة.",
        "إذا تم التحصيل، استخدم تسجيل دفعة من بطاقة العميل أو العقد الصحيح وتأكد من المبلغ قبل الحفظ.",
        "إذا احتاج العميل متابعة لاحقة، استخدم جدولة موعد وحدد تاريخًا واضحًا للرجوع إليه.",
        "إذا استوفى العقد شروط التصعيد، استخدم خيار قانونية بعد مراجعة الملاحظات والدفعات.",
        "في نهاية اليوم، راجع المهام المتبقية ثم صدّر التقرير الشامل عند الحاجة.",
    ]
    for i, step in enumerate(daily_steps, start=1):
        add_number(doc, i, step)

    add_heading(doc, "5. شرح التبويبات")
    add_heading(doc, "5.1 نظرة عامة", 2)
    add_para(doc, "هذا التبويب هو نقطة البداية. يعرض الأولويات، مهام اليوم، مهام التحقق، ملخص الأداء، وآخر الأنشطة.")
    add_bullet(doc, "استخدم قسم يحتاج اهتمامك الفوري للبدء بالعقود ذات المخاطر أو المستحقات الأعلى.")
    add_bullet(doc, "اضغط إنجاز فقط بعد تنفيذ المهمة فعليًا وتوثيق النتيجة المناسبة.")
    add_bullet(doc, "راجع آخر الأنشطة للتأكد من أن إجراءاتك تم تسجيلها بشكل صحيح.")

    add_heading(doc, "5.2 التحصيل الشهري", 2)
    add_para(doc, "هذا التبويب مخصص لمتابعة الفواتير والمبالغ المستحقة حسب العميل أو العقد.")
    add_bullet(doc, "راجع إجمالي المستحق، المحصل، والمتبقي لمعرفة تقدمك خلال الشهر.")
    add_bullet(doc, "افتح مجموعة العميل لمشاهدة الفواتير المرتبطة قبل تسجيل أي دفعة.")
    add_bullet(doc, "لا تسجل دفعة على عميل مشابه بالاسم. تحقق من رقم العقد أو الهاتف قبل الحفظ.", fill=RISK_FILL)
    add_figure(
        doc,
        SCREENSHOTS / "03-collections-tab.png",
        "صورة 5: تبويب التحصيل الشهري من النظام الحي",
        "استخدم هذا التبويب لمراجعة المبالغ المحصلة والمتبقية قبل تسجيل أي دفعة أو متابعة.",
    )

    add_heading(doc, "5.3 العقود", 2)
    add_para(doc, "هذا التبويب يعرض سجل العقود المعيّنة لك، مع البحث والإجراءات الخاصة بكل عقد.")
    add_bullet(doc, "استخدم البحث برقم العقد أو اسم العميل للوصول بسرعة إلى العقد المطلوب.")
    add_bullet(doc, "أزرار الاتصال، الدفعة، الملاحظة، والمتابعة موجودة لتوثيق العمل على العقد مباشرة.")
    add_bullet(doc, "زر قانونية يظهر عند وجود مستحقات ويتطلب مراجعة الحالة قبل التصعيد.")
    add_bullet(doc, "لا يمكن للموظف إلغاء تعيين عقده من مساحة العمل، حتى لو كان لديه دور إداري. الإجراء يتم من إدارة الفريق.", fill=CAUTION_FILL)
    add_figure(
        doc,
        SCREENSHOTS / "04-contracts-tab.png",
        "صورة 6: تبويب العقود والبحث من النظام الحي",
        "من هنا يتم البحث برقم العقد أو اسم العميل، ثم تنفيذ الإجراء المناسب على العقد المحدد.",
    )
    add_bullet(
        doc,
        "إذا لم تظهر عقود في تبويب العقود فهذا يعني أن الحساب الحالي لا توجد عليه عقود نشطة معيّنة في النظام الحي وقت فتح الصفحة.",
        fill=CAUTION_FILL,
    )

    add_heading(doc, "5.4 المهام", 2)
    add_para(doc, "هذا التبويب يجمع المهام الخاصة بك ويعرض حالتها سواء كانت مكتملة أو قيد الانتظار.")
    add_bullet(doc, "ابدأ بالمهام ذات التاريخ الأقرب أو المرتبطة بمبالغ أعلى.")
    add_bullet(doc, "بعد الانتهاء، اضغط إنجاز حتى لا تبقى المهمة مفتوحة في التقارير.")
    add_bullet(doc, "إذا تعذر تنفيذ المهمة، أضف ملاحظة واضحة بسبب التأجيل أو عدم الرد.")
    add_figure(
        doc,
        SCREENSHOTS / "05-tasks-tab.png",
        "صورة 7: تبويب المهام من النظام الحي",
        "يعرض هذا التبويب المهام المفتوحة والمكتملة، ويستخدم لتحديث حالة العمل اليومي.",
    )

    add_heading(doc, "6. شرح الأزرار والإجراءات")
    add_heading(doc, "6.1 أزرار الشريط العلوي", 2)
    add_label_detail_table(
        doc,
        [
            ("الرئيسية", "يعيدك إلى لوحة التحكم الرئيسية. استخدمه إذا أردت الرجوع إلى بطاقات النظام أو اختيار قسم آخر."),
            ("تحديث", "يعيد تحميل بيانات مساحة العمل. استخدمه بعد تسجيل دفعة أو إنهاء مهمة أو عند الشك أن البيانات لم تتحدث."),
            ("تصدير تقرير شامل (Excel)", "ينشئ تقريرًا شاملًا عن أعمالك داخل مساحة العمل. قبل التصدير اضغط تحديث حتى يكون التقرير مبنيًا على آخر بيانات."),
            ("التنبيهات", "يعرض التنبيهات المرتبطة بالمهام أو المتابعات. راجعه في بداية اليوم وخلال العمل إذا ظهرت إشعارات جديدة."),
        ],
    )

    add_heading(doc, "6.2 زر تسجيل مكالمة", 2)
    add_para(doc, "يستخدم هذا الزر لتوثيق أي اتصال مع العميل، سواء تم الرد أو لم يتم الرد. لا يكفي الاتصال من الهاتف بدون تسجيل النتيجة في النظام.")
    add_number(doc, 1, "اضغط تسجيل مكالمة من أزرار الإجراءات السريعة.")
    add_number(doc, 2, "اختر العقد أو العميل الصحيح من القائمة.")
    add_number(doc, 3, "حدد نوع المكالمة: صادرة أو واردة.")
    add_number(doc, 4, "حدد نتيجة المكالمة مثل تم الرد، لا يوجد رد، وعد بالدفع، أو نتيجة مناسبة.")
    add_number(doc, 5, "اكتب ملاحظة مختصرة توضّح ما تم الاتفاق عليه أو سبب عدم الإنجاز.")
    add_number(doc, 6, "إذا كان العميل يحتاج متابعة لاحقة، فعّل خيار يتطلب متابعة لاحقة ثم احفظ المكالمة.")
    add_bullet(doc, "القاعدة المهمة: كل مكالمة مؤثرة على التحصيل يجب أن تكون موثقة بنتيجة واضحة.", fill=SUCCESS_FILL)
    add_figure(
        doc,
        SCREENSHOTS / "08-call-dialog.png",
        "صورة 8: نافذة تسجيل مكالمة",
        "تظهر فيها حقول اختيار العقد، نوع المكالمة، نتيجة المكالمة، الغرض، مدة المكالمة، والملاحظات.",
        width_in=4.2,
    )

    add_heading(doc, "6.3 زر تسجيل دفعة", 2)
    add_para(doc, "يستخدم هذا الزر لتسجيل دفعة على العميل أو العقد الصحيح. إذا لم يكن هناك عميل محدد، سيطلب منك النظام اختيار العميل أو الضغط من بطاقة العقد المطلوبة.")
    add_number(doc, 1, "ابدأ من بطاقة العميل أو العقد عندما تكون متأكدًا من هوية العميل.")
    add_number(doc, 2, "اضغط تسجيل دفعة.")
    add_number(doc, 3, "راجع رقم العقد واسم العميل قبل إدخال المبلغ.")
    add_number(doc, 4, "أدخل مبلغ الدفعة وطريقة الدفع وأي مرجع مطلوب.")
    add_number(doc, 5, "راجع البيانات مرة أخيرة ثم احفظ الدفعة.")
    add_bullet(doc, "لا تسجل دفعة من الصفحة العامة إذا لم تكن اخترت العميل أو العقد؛ ابحث أولًا عن العقد الصحيح ثم سجّل الدفعة منه.", fill=RISK_FILL)
    add_figure(
        doc,
        SCREENSHOTS / "09-payment-action.png",
        "صورة 9: تنبيه تسجيل الدفعة بدون اختيار عميل",
        "عند الضغط على تسجيل دفعة بدون تحديد عقد، يوضح النظام أنه يجب اختيار العميل أو العقد المطلوب أولًا.",
    )

    add_heading(doc, "6.4 زر جدولة موعد", 2)
    add_para(doc, "يستخدم هذا الزر لإنشاء موعد متابعة قادم، مثل متابعة دفعة شهرية، وعد دفع، أو مراجعة حالة عميل.")
    add_number(doc, 1, "اضغط جدولة موعد.")
    add_number(doc, 2, "اختر العقد أو العميل الصحيح.")
    add_number(doc, 3, "حدد نوع المتابعة، ثم اكتب عنوانًا واضحًا مثل متابعة دفعة شهر يناير.")
    add_number(doc, 4, "اختر التاريخ والوقت، وحدد الأولوية حسب أهمية الحالة.")
    add_number(doc, 5, "اكتب ملاحظة اختيارية إذا كانت هناك تفاصيل تساعدك عند الرجوع للموعد.")
    add_number(doc, 6, "اضغط جدولة المتابعة.")
    add_bullet(doc, "استخدم المواعيد بدل الاعتماد على الذاكرة، خصوصًا في وعود الدفع أو الحالات المتكررة.", fill=SUCCESS_FILL)
    add_figure(
        doc,
        SCREENSHOTS / "10-schedule-dialog.png",
        "صورة 10: نافذة جدولة متابعة",
        "توضح حقول العقد، نوع المتابعة، العنوان، التاريخ، الوقت، الأولوية، والملاحظات.",
        width_in=4.2,
    )

    add_heading(doc, "6.5 زر ملاحظة جديدة", 2)
    add_para(doc, "يستخدم هذا الزر لتوثيق معلومة مهمة عن العميل أو العقد لا تعتبر دفعة أو مكالمة، مثل طلب مراجعة، ملاحظة إدارية، أو معلومة من العميل.")
    add_number(doc, 1, "اضغط ملاحظة جديدة.")
    add_number(doc, 2, "اختر العقد أو العميل الصحيح.")
    add_number(doc, 3, "حدد نوع الملاحظة.")
    add_number(doc, 4, "اكتب محتوى الملاحظة بوضوح وباختصار.")
    add_number(doc, 5, "إذا كانت الملاحظة مهمة وتحتاج ظهورًا أعلى، فعّل خيار ملاحظة مهمة.")
    add_number(doc, 6, "اضغط حفظ الملاحظة.")
    add_bullet(doc, "الملاحظة الجيدة تجيب على سؤال: ماذا حدث؟ وماذا يجب أن يفعل الموظف التالي؟", fill=CAUTION_FILL)
    add_figure(
        doc,
        SCREENSHOTS / "11-note-dialog.png",
        "صورة 11: نافذة إضافة ملاحظة",
        "توضح اختيار العقد، نوع الملاحظة، محتوى الملاحظة، وخيار تمييزها كملاحظة مهمة.",
        width_in=4.2,
    )

    add_heading(doc, "6.6 زر قانونية داخل بطاقة العقد", 2)
    add_para(doc, "يظهر هذا الإجراء عند وجود عقد عليه مستحقات ويحتاج تصعيدًا. لا يستخدم للتحويل العشوائي، بل بعد مراجعة سجل الدفعات والملاحظات ومحاولات التواصل.")
    add_number(doc, 1, "افتح تبويب العقود وابحث عن العقد برقم العقد أو اسم العميل.")
    add_number(doc, 2, "راجع المبلغ المستحق وآخر ملاحظة وآخر متابعة.")
    add_number(doc, 3, "إذا كانت الحالة تستحق التصعيد، اضغط قانونية.")
    add_number(doc, 4, "اكتب سبب التحويل بوضوح، مثل عدم تجاوب متكرر أو وعد دفع غير منفذ.")
    add_number(doc, 5, "بعد التحويل، تابع تعليمات الشؤون القانونية ولا تغير حالة العقد من نفسك.")

    add_heading(doc, "7. متى يتم التحويل إلى الشؤون القانونية؟")
    add_para(doc, "التحويل القانوني ليس إجراءً عاديًا لكل تأخير. استخدمه عندما تكون بيانات العقد واضحة، والمستحقات مؤكدة، ومحاولات المتابعة موثقة.")
    add_bullet(doc, "تأكد من رقم العقد واسم العميل والمبلغ المستحق قبل التحويل.")
    add_bullet(doc, "راجع آخر الدفعات والملاحظات حتى لا يتم تحويل عقد تم سداده أو توجد عليه تسوية قريبة.")
    add_bullet(doc, "اكتب سببًا واضحًا للتصعيد مثل: عدم تجاوب متكرر، وعد دفع غير منفذ، مبلغ كبير متأخر، أو رفض السداد.")
    add_bullet(doc, "بعد التحويل، تابع حسب تعليمات الإدارة القانونية ولا تعدل حالة العقد بدون توجيه.")

    add_heading(doc, "8. قواعد مهمة للموظفين")
    add_bullet(doc, "لا تستخدم حساب موظف آخر ولا تسجل إجراءات باسم شخص آخر.")
    add_bullet(doc, "لا تحفظ دفعة قبل مطابقة العميل والعقد والمبلغ وطريقة الدفع.")
    add_bullet(doc, "لا تترك المكالمات المهمة بدون ملاحظة أو نتيجة واضحة.")
    add_bullet(doc, "لا تلغي تعيين عقدك من مساحة العمل. إذا كان هناك خطأ في التعيين، أبلغ مدير الفريق ليتم التصحيح من إدارة الفريق.")
    add_bullet(doc, "استخدم لغة مهنية مختصرة في الملاحظات، وتجنب العبارات غير الواضحة مثل تم التواصل فقط بدون نتيجة.")

    add_heading(doc, "9. سيناريوهات شائعة")
    add_matrix_table(
        doc,
        ["الحالة", "الإجراء الصحيح", "ملاحظة"],
        [
            ["العميل دفع مبلغًا", "افتح العقد الصحيح ثم سجّل الدفعة", "راجع المبلغ وطريقة الدفع قبل الحفظ"],
            ["العميل لا يرد", "سجّل مكالمة بنتيجة لا يوجد رد ثم جدولة متابعة", "لا تكرر الاتصال بدون توثيق"],
            ["وعد بالدفع غدًا", "سجّل المكالمة ثم أنشئ موعد متابعة", "اكتب تاريخ الوعد بوضوح"],
            ["العقد متأخر بمبلغ كبير", "راجع الدفعات والملاحظات ثم قيّم التحويل للقانونية", "التصعيد يحتاج سببًا موثقًا"],
            ["العقد معيّن للموظف الخطأ", "أبلغ مدير الفريق", "إعادة التعيين تتم من إدارة الفريق"],
            ["نهاية اليوم", "راجع المهام المتبقية وصدّر التقرير عند الحاجة", "استخدم تحديث قبل التصدير"],
        ],
        [2200, 3860, 3300],
    )

    add_heading(doc, "10. قائمة تحقق نهاية اليوم")
    checklist = [
        "تمت مراجعة العملاء ذوي الأولوية.",
        "تم تسجيل نتائج المكالمات المهمة.",
        "تم تسجيل الدفعات على العقود الصحيحة.",
        "تمت جدولة المتابعات المؤجلة.",
        "تم إنجاز المهام المكتملة فعليًا فقط.",
        "تم توثيق الحالات التي تحتاج مراجعة مدير الفريق.",
        "تم تحديث الصفحة قبل استخراج أي تقرير.",
    ]
    for item in checklist:
        add_bullet(doc, f"□ {item}")

    add_heading(doc, "11. عند وجود مشكلة")
    add_bullet(doc, "إذا لم تظهر بياناتك، اضغط تحديث وتأكد من اتصال الإنترنت ثم أعد فتح الصفحة.")
    add_bullet(doc, "إذا وجدت عقدًا غير تابع لك، أبلغ مدير الفريق بدل محاولة إلغاء التعيين من مساحة العمل.")
    add_bullet(doc, "إذا ظهر خطأ في مبلغ أو اسم عميل، لا تكمل الإجراء قبل مراجعة الإدارة أو المسؤول المختص.")
    add_bullet(doc, "إذا لم تتمكن من فتح الصفحة، التقط صورة للخطأ وأرسلها للدعم الفني مع وقت حدوث المشكلة.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
