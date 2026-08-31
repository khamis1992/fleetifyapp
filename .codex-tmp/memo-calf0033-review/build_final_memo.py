from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\khamis\Desktop\مذكرة_شارحة_معد_لة_CALF0033.docx")
OUTPUT = Path(r"C:\Users\khamis\Documents\fleetifyapp\artifacts\مذكرة_شارحة_منقحة_C-ALF-0033.docx")
LRM = "\u200e"


def set_rtl_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15


def format_run(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_text(paragraph, text, *, bold_prefix=None, size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    clear_paragraph(paragraph)
    set_rtl_paragraph(paragraph, alignment)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        format_run(first, size=size, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        format_run(rest, size=size)
    else:
        run = paragraph.add_run(text)
        format_run(run, size=size)


def set_heading(paragraph, text, level=1):
    clear_paragraph(paragraph)
    set_rtl_paragraph(paragraph)
    run = paragraph.add_run(text)
    format_run(run, size=14 if level == 1 else 12, bold=True, color=(22, 78, 99))
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True


def delete_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def insert_before(reference, text, *, bold_prefix=None):
    new_p = OxmlElement("w:p")
    reference._p.addprevious(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    set_text(paragraph, text, bold_prefix=bold_prefix)
    return paragraph


def set_cell_text(cell, text, *, bold=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_rtl_paragraph(paragraph, alignment)
    run = paragraph.add_run(text)
    format_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def set_cell_lines(cell, lines, *, size=10.5):
    cell.text = ""
    for index, (label, value) in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        clear_paragraph(paragraph)
        set_rtl_paragraph(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
        paragraph.paragraph_format.space_after = Pt(1)
        label_run = paragraph.add_run(label)
        format_run(label_run, size=size, bold=True)
        value_run = paragraph.add_run(value)
        format_run(value_run, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_simple_field(paragraph, instruction):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "666666")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    rpr.extend([color, size])
    text = OxmlElement("w:t")
    text.text = "1"
    run_el.extend([rpr, text])
    fld.append(run_el)
    paragraph._p.append(fld)


def add_footer(section):
    section.footer_distance = Cm(0.7)
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    set_rtl_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "3")
    top.set(qn("w:color"), "B8C2CC")
    pbdr.append(top)

    r1 = paragraph.add_run("شركة العراف لتأجير السيارات ذ.م.م | السجل التجاري 146832 | الدوحة - قطر | khamis-1992@hotmail.com | صفحة ")
    format_run(r1, size=8, color=(102, 102, 102))
    add_simple_field(paragraph, "PAGE")
    r2 = paragraph.add_run(" من ")
    format_run(r2, size=8, color=(102, 102, 102))
    add_simple_field(paragraph, "NUMPAGES")


def remove_internal_appendix(document):
    marker = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("ملحق داخلي"):
            marker = paragraph._p
            break
    if marker is None:
        return
    body = document._body._element
    deleting = False
    for child in list(body):
        if child is marker:
            deleting = True
        if deleting and child.tag != qn("w:sectPr"):
            body.remove(child)


def replace_by_start(document, start, text, *, heading=False, level=1, bold_prefix=None):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(start):
            if heading:
                set_heading(paragraph, text, level=level)
            else:
                set_text(paragraph, text, bold_prefix=bold_prefix)
            return paragraph
    raise RuntimeError(f"Paragraph not found: {start}")


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    remove_internal_appendix(doc)

    # الترويسة
    set_text(doc.paragraphs[0], "مذكرة شارحة", size=17, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.paragraphs[0].runs[0].bold = True
    set_text(
        doc.paragraphs[1],
        "مقدمة إلى محكمة الاستثمار والتجارة – الدائرة الابتدائية",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.paragraphs[1].runs[0].bold = True

    # بيانات الأطراف والعقد
    info = doc.tables[0]
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = True
    info_rows = [
        ("شركة العراف لتأجير السيارات ذ.م.م — السجل التجاري 146832", "المدعية:"),
        ("عبد الرحيم شاكر أحمد محمد — الرقم الشخصي 28373601770 — الجنسية: السودان — العنوان: الدوحة، قطر", "المدعى عليه:"),
        (f"رقم {LRM}C-ALF-0033{LRM} — بداية 08/02/2024 — نهاية العقد 08/01/2027 — أجرة شهرية 1,700 ريال قطري", "عقد الإيجار محل النزاع:"),
        (f"نوع المركبة: ({LRM}GAC GS3{LRM}) — الموديل: 2024 — اللوحة: (4018) — رقم الهيكل: ({LRM}LMWHR1G26R1107033{LRM})", "المركبة:"),
        ("فسخ عقد إيجار مركبة", "موضوع الدعوى:"),
    ]
    for row, (value, label) in zip(info.rows, info_rows):
        set_cell_text(row.cells[0], value)
        set_cell_text(row.cells[1], label, bold=True)
        shade_cell(row.cells[1], "EAF2F5")
    set_cell_lines(
        info.rows[3].cells[0],
        [
            ("نوع المركبة: ", f"{LRM}GAC GS3{LRM}"),
            ("الموديل: ", "2024"),
            ("رقم اللوحة: ", "4018"),
            ("رقم الهيكل: ", f"{LRM}LMWHR1G26R1107033{LRM}"),
        ],
    )

    # الاختصاص والوقائع
    replace_by_start(doc, "أولاً:", "أولاً: الاختصاص القضائي", heading=True)
    replace_by_start(
        doc,
        "تختص محكمة الاستثمار",
        "تختص محكمة الاستثمار والتجارة بنظر هذه الدعوى عملاً بالمادة (7) من قانون رقم (21) لسنة 2021 بإصدار قانون إنشاء محكمة الاستثمار والتجارة، باعتبار النزاع ناشئاً عن عقد تجاري يتعلق بتأجير مركبة، وينعقد الاختصاص المكاني لمحاكم دولة قطر وفقاً للعقد والقواعد المقررة قانوناً.",
    )
    replace_by_start(doc, "ثانياً:", "ثانياً: الوقائع", heading=True)
    replace_by_start(
        doc,
        "أبرمت المدعية",
        f"أبرمت المدعية مع المدعى عليه عقد إيجار المركبة المبينة بياناتها أعلاه رقم ({LRM}C-ALF-0033{LRM}) بتاريخ 08/02/2024، لمدة محددة تنتهي بتاريخ 08/01/2027، وبأجرة شهرية مقدارها (1,700) ريال قطري، وذلك وفق عقد الإيجار المرفق.",
    )
    replace_by_start(
        doc,
        "تتعلق العلاقة الإيجارية",
        "تثبت العلاقة الإيجارية وبيانات المركبة والتزامات الطرفين بعقد الإيجار والمستندات المرفقة، وتخضع واقعة الحيازة والرد لما يثبت انتقال الحيازة فعلياً.",
    )
    replace_by_start(
        doc,
        "إلا أن المدعى عليه",
        "إلا أن المدعى عليه أخل بالتزامه بسداد الأجرة، إذ تخلف عن سداد الفواتير المبينة تفصيلاً بكشف الحساب المرفق، والمستحقة خلال الفترة من 01/01/2025 حتى 01/08/2026.",
    )
    replace_by_start(
        doc,
        "بلغ إجمالي",
        "بلغ إجمالي الفواتير محل المطالبة مبلغ (19,200) ريال قطري، سدد منه المدعى عليه مبلغ (5,160) ريال قطري، فأصبح صافي الإيجارات غير المسددة حتى 01/08/2026 مبلغ (14,040) ريال قطري، وفق كشف الحساب والفواتير وإيصالات السداد المرفقة.",
    )
    # حذف الادعاء غير المؤيد بأن المركبة ما زالت في حيازته.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("وعلى الرغم مما تقدم"):
            delete_paragraph(paragraph)

    # البيان الحسابي
    replace_by_start(doc, "ثالثاً:", "ثالثاً: البيان الحسابي للمطالبة", heading=True)
    claims = doc.tables[1]
    claims.alignment = WD_TABLE_ALIGNMENT.CENTER
    claim_rows = [
        ("المبلغ (ريال قطري)", "الفترة / المستند", "البند"),
        ("19,200", "الفواتير المبينة بكشف الحساب عن الفترة من 01/01/2025 حتى 01/08/2026", "إجمالي الفواتير محل المطالبة"),
        ("(5,160)", "وفق إيصالات السداد وكشف الحساب", "يُخصم: المبالغ المسددة"),
        ("14,040", "حتى 01/08/2026", "صافي الإيجارات غير المسددة"),
        ("14,040", "وفق كشف المطالبة المرفق", "صافي المطالبة حتى تاريخ الكشف"),
    ]
    for r_index, (amount, period, item) in enumerate(claim_rows):
        row = claims.rows[r_index]
        for c_index, value in enumerate((amount, period, item)):
            set_cell_text(
                row.cells[c_index],
                value,
                bold=(r_index == 0 or r_index == len(claim_rows) - 1),
                alignment=WD_ALIGN_PARAGRAPH.CENTER if c_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT,
                size=9.5,
            )
            if r_index == 0:
                shade_cell(row.cells[c_index], "DCEAF0")
            elif r_index == len(claim_rows) - 1:
                shade_cell(row.cells[c_index], "E8F4EF")
    replace_by_start(
        doc,
        "ويؤكد المدعية",
        "وتؤكد المدعية أن البيان السابق لا يتضمن ازدواجاً في المطالبة؛ فلا تجمع عن المركبة والفترة الزمنية ذاتها بين الأجرة التعاقدية وتعويض الاحتباس، ولا يتكرر أصل الدين ضمن أي عنصر من عناصر الضرر، وكل مبلغ خاضع لما يؤيده من عقد أو فاتورة أو كشف رسمي أو إيصال.",
    )

    # الأساس القانوني
    replace_by_start(doc, "رابعاً:", "رابعاً: الأساس القانوني", heading=True)
    replace_by_start(
        doc,
        "تستند هذه الدعوى",
        "تستند هذه الدعوى إلى أحكام القانون المدني القطري رقم (22) لسنة 2004، وقانون إنشاء محكمة الاستثمار والتجارة رقم (21) لسنة 2021، وقانون التنفيذ القضائي رقم (4) لسنة 2024، وذلك على النحو الآتي:",
    )
    for title in (
        ("1. القوة الملزمة", "١. القوة الملزمة للعقد وحسن النية — المادتان ١٧١ و١٧٢"),
        ("2. الالتزام", "٢. الالتزام بسداد الأجرة — المادة ٦٠٧"),
        ("3. الفسخ", "٣. الفسخ القضائي — المادة ١٨٣"),
        ("4. رد المركبة", "٤. رد المركبة وتعويض التأخر في ردها — المواد ٦١٦ و٦١٧ و٦١٨"),
        ("5. التعويض", "٥. تعويض الاحتباس ومنع الازدواج"),
        ("6. الشرط", "٦. التعويض الاتفاقي والتعويض عن التأخر في الوفاء — المواد ٢٥٦ و٢٦٣ و٢٦٦ و٢٦٧ و٢٦٨"),
    ):
        replace_by_start(doc, title[0], title[1], heading=True, level=2)
    replace_by_start(
        doc,
        "تقرر المادة (171)",
        "تقرر المادة (171) أن العقد شريعة المتعاقدين، وتوجب المادة (172) تنفيذه طبقاً لما اشتمل عليه وبطريقة تتفق مع حسن النية. والثابت من المستندات قيام العلاقة العقدية، في حين تخلف المدعى عليه عن الوفاء بكامل الأجرة المستحقة.",
    )
    replace_by_start(
        doc,
        "تجيز المادة (183)",
        "تجيز المادة (183) طلب فسخ العقد عند عدم وفاء أحد المتعاقدين بالتزامه، مع التعويض إن كان له مقتضى، وذلك بعد الإعذار على الوجه المقرر قانوناً. وإذ يتعلق الإخلال بالتزام جوهري ومتجدد هو سداد الأجرة، تتمسك المدعية بطلب الفسخ وترتيب آثاره من التاريخ الذي تحدده المحكمة.",
    )
    replace_by_start(
        doc,
        "توجب المادة (616)",
        "تقضي المادة (616) بأن المستأجر يلتزم برد العين المؤجرة عند انتهاء الإيجار، فإذا أبقاها تحت يده دون حق التزم بأن يدفع للمؤجر تعويضاً يراعى في تقديره القيمة الإيجارية وما أصاب المؤجر من ضرر، مع مراعاة أحكام المادتين (617) و(618) بشأن حالة الرد ومصروفاته.",
    )
    replace_by_start(
        doc,
        "إذا استمر المستأجر",
        "إذا ثبت استمرار حيازة المدعى عليه للمركبة بعد صيرورة الفسخ منتجاً لآثاره، استحق عن مدة الاحتباس تعويض يقدر على أساس القيمة الإيجارية السوقية وما يثبت من ضرر إضافي، دون الجمع بين الأجرة التعاقدية وتعويض الاحتباس عن المدة ذاتها.",
    )
    replace_by_start(
        doc,
        "تجيز المادة (265)",
        "تنص المادة (4) من عقد الإيجار، بحسب العقد المرفق، على تعويض اتفاقي بواقع (1,200) ريال قطري عن كل شهر تأخير. وتتمسك المدعية بإعماله في الحدود التي يجيزها القانون، وبالقدر المنطبق على واقعة التأخير محل الدعوى، ودون ازدواج مع تعويض آخر عن الضرر ذاته.",
    )
    replace_by_start(
        doc,
        "مع مراعاة أن المادة (266)",
        "وتخضع قيمة التعويض الاتفاقي لرقابة المحكمة وفق المادتين (266) و(267)، بما في ذلك سلطة تخفيضه إذا لم يقع ضرر أو كان التقدير مبالغاً فيه، وعدم مجاوزته إلا إذا ثبت الغش أو الخطأ الجسيم.",
    )
    replace_by_start(
        doc,
        "كما تجيز المادة (267)",
        "أما الضرر الناشئ عن التأخر في الوفاء بالدين النقدي فيخضع للمادتين (256) و(268)، ويشترط للحكم به ثبوت الإعذار والضرر وعلاقة السببية، وتدخل الخسارة وما فات من كسب في نطاق التعويض وفق المادة (263)، مع منع تكرار التعويض عن الضرر ذاته.",
    )
    replace_by_start(doc, "الدفوع المتوقعة", "خامساً: الإثبات والرد على الدفوع", heading=True)
    replace_by_start(
        doc,
        "يقع على عاتق المدعى عليه",
        "تلتزم المدعية بإثبات العقد والدين وعناصر الضرر بالمستندات، بينما يقع على المدعى عليه إثبات ما يدعيه من سداد أو رد فعلي للمركبة أو انتقال للحيازة. وتحتفظ المدعية بحقها في تقديم ما يستجد من كشوف وإيصالات ومحاضر رسمية رداً على أي دفع يثار أثناء نظر الدعوى.",
    )

    # الطلبات: إعادة بنائها بالكامل دون الطلب الاحتياطي المتعلق باعتبار الإعلان إعذاراً.
    replace_by_start(doc, "خامساً: الطلبات", "سادساً: الطلبات", heading=True)
    intro = replace_by_start(doc, "لذلك،", "لذلك، تلتمس المدعية من المحكمة الموقرة الحكم بما يلي:")
    respect = next(p for p in doc.paragraphs if p.text.strip().startswith("وتفضلوا"))
    current = False
    for paragraph in list(doc.paragraphs):
        if paragraph._p is intro._p:
            current = True
            continue
        if paragraph._p is respect._p:
            break
        if current:
            delete_paragraph(paragraph)

    requests = [
        ("أولاً:", "أولاً: قبول الدعوى شكلاً."),
        ("ثانياً:", f"ثانياً: الحكم بفسخ عقد إيجار المركبة رقم ({LRM}C-ALF-0033{LRM}) لإخلال المدعى عليه إخلالاً جوهرياً ومستمراً بالتزامه بسداد الأجرة، وترتيب آثار الفسخ من التاريخ الذي تحدده المحكمة."),
        ("ثالثاً:", "ثالثاً: إلزام المدعى عليه بأن يؤدي للمدعية مبلغ (14,040) ريال قطري قيمة صافي الأجرة المستحقة حتى 01/08/2026، وما يستجد من أجرة تعاقدية بواقع (1,700) ريال قطري شهرياً، أو نسبتها عن جزء الشهر وفق العقد، حتى التاريخ الذي يصير فيه الفسخ منتجاً لآثاره، بعد خصم أي مبالغ مسددة."),
        ("رابعاً:", "رابعاً: إلزام المدعى عليه برد المركبة المبينة بياناتها في صدر هذه المذكرة، تسليماً فعلياً كاملاً وصالحاً للحيازة والانتفاع، مع جميع المفاتيح والوثائق والملحقات، ولا يعتد بالرد إلا بما يثبت انتقال الحيازة فعلياً، ومن ذلك محضر تسليم أو مستند رسمي معتبر."),
        ("خامساً:", "خامساً: إذا ثبت استمرار حيازة المدعى عليه للمركبة، إلزامه من اليوم التالي لصيرورة الفسخ منتجاً لآثاره وحتى تاريخ التسليم الفعلي بتعويض احتباس يحدد وفق القيمة الإيجارية السوقية للمركبة، مع تعويض ما يثبت من أضرار إضافية، ودون الجمع بين الأجرة والتعويض عن الفترة ذاتها."),
        ("سادساً:", "سادساً: إلزام المدعى عليه بقيمة إصلاح الأضرار غير الناتجة عن الاستعمال المألوف، وقيمة النقص في القيمة السوقية، وقيمة الملحقات والمفاتيح المفقودة، ومصاريف الفحص والسحب والاسترداد والحجز والتأمين، في حدود ما تثبته المستندات أو الخبرة الفنية."),
        ("سابعاً:", "سابعاً: إلزام المدعى عليه بتعويض المدعية عما يثبت من فوات الانتفاع وصافي الكسب خلال مدة إصلاح المركبة المعقولة بعد استردادها، وفق المستندات أو ما تقدره المحكمة، ودون ازدواج مع تعويض الاحتباس."),
        ("ثامناً:", "ثامناً: إلزام المدعى عليه بتعويض عادل عن الضرر الناجم عن التأخر في سداد الدين النقدي بعد إعذاره قانوناً، وفق المواد ٢٥٦ و٢٦٣ و٢٦٨ من القانون المدني، بالمبلغ الذي تثبته المستندات أو تقدره المحكمة، ودون الجمع بينه وبين تعويض اتفاقي عن الضرر ذاته."),
        ("تاسعاً:", "تاسعاً: إلزام المدعى عليه بالتعويض الاتفاقي المنصوص عليه في المادة ٤ من عقد الإيجار بواقع (1,200) ريال قطري عن كل شهر تأخير، في الحدود التي يجيزها القانون وبقدر انطباق البند على واقعة التأخير محل الدعوى، ودون ازدواج مع أي تعويض عن الضرر ذاته."),
        ("عاشراً:", "عاشراً: إلزام المدعى عليه بقيمة المخالفات المرورية والرسوم والمصاريف التي يثبت وقوعها خلال فترة حيازته للمركبة، وفق الكشوف الرسمية وأحكام العقد."),
        ("حادي عشر:", "حادي عشر: في حال تعذر الرد العيني، إلزام المدعى عليه، على سبيل البديل، بالقيمة السوقية للمركبة وقت وجوب ردها، مع التعويضات الأخرى التي لا تتداخل مع قيمة المركبة، وفق المستندات أو تقدير الخبرة عند الاقتضاء."),
        ("ثاني عشر:", "ثاني عشر: شمول الحكم بالنفاذ المعجل، وبغير كفالة إن رأت المحكمة توافر شروط المادة ٩ من قانون التنفيذ القضائي رقم ٤ لسنة ٢٠٢٤، وعلى الأخص ما يترتب على تأخير التنفيذ من ضرر جسيم بمصلحة المدعية."),
        ("ثالث عشر:", "ثالث عشر: إلزام المدعى عليه بالرسوم والمصاريف ومقابل أتعاب المحاماة."),
    ]
    for prefix, text in requests:
        insert_before(respect, text, bold_prefix=prefix)
    set_text(respect, "وتفضلوا بقبول فائق الاحترام والتقدير،،،")

    # تنسيق التوقيع وتثبيت التاريخ فارغاً.
    signature_map = {
        "شركة العراف": "شركة العراف لتأجير السيارات ذ.م.م",
        "خميس هاشم": "خميس هاشم الجبر",
        "المخول بالتوقيع": "المخول بالتوقيع",
        "التوقيع:": "التوقيع: __________________",
        "التاريخ:": "التاريخ: __________________",
    }
    for start, text in signature_map.items():
        replace_by_start(doc, start, text, bold_prefix=start if start in ("التوقيع:", "التاريخ:") else None)

    # إزالة أي تمييز لوني أو نصوص مسودة متبقية وتوحيد الخط.
    for paragraph in doc.paragraphs:
        set_rtl_paragraph(paragraph, paragraph.alignment or WD_ALIGN_PARAGRAPH.RIGHT)
        for run in paragraph.runs:
            run.font.highlight_color = None
            if run.font.size is None:
                format_run(run, size=11, bold=bool(run.bold))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_rtl_paragraph(paragraph, paragraph.alignment or WD_ALIGN_PARAGRAPH.RIGHT)
                    for run in paragraph.runs:
                        run.font.highlight_color = None

    # هوامش وتذييل ثابت على جميع الصفحات.
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        add_footer(section)

    # منع انفصال عناوين الأقسام عن الفقرة التالية.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(("أولاً: الاختصاص", "ثانياً: الوقائع", "ثالثاً: البيان", "رابعاً: الأساس", "خامساً: الإثبات", "سادساً: الطلبات")):
            paragraph.paragraph_format.keep_with_next = True

    # إزالة فاصل الصفحة الذي كان يسبق الملحق الداخلي المحذوف.
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip() and paragraph._p.findall(".//w:br[@w:type='page']", paragraph._p.nsmap):
            delete_paragraph(paragraph)

    doc.core_properties.title = "مذكرة شارحة - C-ALF-0033"
    doc.core_properties.subject = "فسخ عقد إيجار مركبة"
    doc.core_properties.comments = "نسخة منقحة للتقديم؛ لا تتضمن ملاحظات داخلية أو أرقام دعوى."
    doc.save(OUTPUT)
    print("Final memo created")


if __name__ == "__main__":
    main()
